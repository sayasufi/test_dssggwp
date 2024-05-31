import datetime

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django_filters.rest_framework import DjangoFilterBackend
from docx import Document
from rest_framework import viewsets, filters
from rest_framework.decorators import action, api_view
from rest_framework.response import Response

from accounts.forms import RoomBookingForm, BookingFilterForm
from accounts.models import RoomBooking, Room
from accounts.serializers import RoomSerializer, RoomBookingSerializer


def home(request):
    return render(request, 'accounts/home.html')


@login_required
def profile(request):
    return render(request, 'accounts/profile.html')


@login_required
def book_room(request):
    if request.method == 'POST':
        form = RoomBookingForm(request.POST)
        if form.is_valid():
            room = form.cleaned_data['room']
            start_time = form.cleaned_data['start_time']
            end_time = form.cleaned_data['end_time']
            # Check for conflicts
            if RoomBooking.objects.filter(room=room, start_time__lt=end_time, end_time__gt=start_time).exists():
                form.add_error(None, 'Комната уже забронирована на это время')
            else:
                booking = form.save(commit=False)
                booking.user = request.user
                booking.save()
                return redirect('booking_list')
    else:
        form = RoomBookingForm()
    return render(request, 'accounts/book_room.html', {'form': form})


@login_required
def booking_list(request):
    form = BookingFilterForm(request.GET or None)
    bookings = RoomBooking.objects.all()

    if form.is_valid():
        if form.cleaned_data['room']:
            bookings = bookings.filter(room=form.cleaned_data['room'])
        if form.cleaned_data['start_date']:
            bookings = bookings.filter(start_time__gte=form.cleaned_data['start_date'])
        if form.cleaned_data['end_date']:
            bookings = bookings.filter(end_time__lte=form.cleaned_data['end_date'])

    return render(request, 'accounts/booking_list.html', {'form': form, 'bookings': bookings})


class RoomViewSet(viewsets.ModelViewSet):
    queryset = Room.objects.all()
    serializer_class = RoomSerializer


class RoomBookingViewSet(viewsets.ModelViewSet):
    queryset = RoomBooking.objects.all()
    serializer_class = RoomBookingSerializer
    filter_backends = [DjangoFilterBackend, filters.OrderingFilter]
    filterset_fields = ['room', 'start_time', 'end_time']
    ordering_fields = ['start_time', 'end_time']

    @action(detail=False, methods=['get'])
    def by_date_range(self, request):
        start_date = request.query_params.get('start_date')
        end_date = request.query_params.get('end_date')
        room_id = request.query_params.get('room_id')

        bookings = RoomBooking.objects.all()

        if start_date:
            bookings = bookings.filter(start_time__gte=start_date)
        if end_date:
            bookings = bookings.filter(end_time__lte=end_date)
        if room_id:
            bookings = bookings.filter(room_id=room_id)

        serializer = self.get_serializer(bookings, many=True)
        return Response(serializer.data)

@api_view(['GET'])
def generate_report(request):
    start_date = request.query_params.get('start_date')
    end_date = request.query_params.get('end_date')
    room_id = request.query_params.get('room_id')

    bookings = RoomBooking.objects.all()

    if start_date:
        bookings = bookings.filter(start_time__gte=start_date)
    if end_date:
        bookings = bookings.filter(end_time__lte=end_date)
    if room_id:
        bookings = bookings.filter(room_id=room_id)

    # Создание документа
    doc = Document()
    doc.add_heading('Отчет о бронированиях', 0)

    for booking in bookings:
        doc.add_paragraph(
            f"Комната: {booking.room.room_number}\n"
            f"Начало: {booking.start_time}\n"
            f"Окончание: {booking.end_time}\n"
            f"Цель: {booking.purpose}\n"
            f"Пользователь: {booking.user.username}\n"
        )

    # Настройка имени файла
    report_filename = f"booking_report_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

    # Создание HTTP ответа
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={report_filename}'
    doc.save(response)
    return response